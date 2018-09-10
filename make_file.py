from docx import Document           #doc import
from docxpl import DocxTemplate

import os                            #os import

class CopyPic:

    filepath = ""    #폴더 위치
    size = 0                #라벨 크기 선택
    size_check = 0          #사진 갯수 확인

    table_width = 0         #그림 너비
    table_height = 0        #그림 높이

    makedfile_list = []

    def rec_filepath(self, filepath):
        self.filepath = filepath


    def rec_size(self, size) :
        self.size = size
        if size != 24 and size != 21 :
            print("잘못된 입력입니다")
            return 0
        return 1


    def excute_copy(self) :
        # 폴더 내 파일 이름 불러오기
        path_dir = self.filepath
        file_list = os.listdir(path_dir)
        file_list.sort()

        length = len(file_list)             #그림 파일 갯수

        self.size_check = length//self.size + 1  #그림 갯수를 확인하여 저장 파일 정함

        print("길이는 %d" % self.size_check)


        for output in range(self.size_check) :

            document = Document(str(self.size) + "labels.docx")

            table = document.tables  # table 생성

            self.table_height = table[0].rows[0].height
            self.table_width = table[0].rows[0].cells[0].width

            print(self.table_height)

            # 초기값 설정

            row = 0
            col = 0

            for name in file_list[output*self.size : output*self.size + self.size]:
                p = table[0].rows[row].cells[col].paragraphs
                r = p[0].add_run()  # 작성?
                r.add_picture(self.filepath + name, width=(self.table_width-100000), height=(self.table_height-100000))

                # 한칸씩 이동
                col += 1

                if row == self.size // 3:
                    row = 0
                if col == 3:
                    col = 0
                    row = row + 1

                document.save(str(self.size) + 'labels-'+ str(output) + '.docx')

            self.makedfile_list.append(str(self.size) + 'labels-' + str(output) + '.docx')


class merge_docs():
    def __init__(self,*args, save_as = 'dobble'):
        self.docs = []
        self.subdoc_nums = []
        for i , arg in enumerate(args):
            self.docs.append(arg)
            self.subdoc_nums.append('subdoc{}'.format(i))
        self.merge(save_as)

    def merge(self, save_as):
        tpl=DocxTemplate('merge_template.docx')
        new_subdocs = []
        for doc in self.docs :
            new_subdocs.append(tpl.new_subdoc(doc))
        context = {
                subdoc_num : subdoc for subdoc_num, subdoc in zip(self.subdoc_nums, new_subdocs)
        }
        tpl.render(context)
        tpl.save(save_as + '.docx')

        def unlink_docs(self) :
            for doc in self.docs :
                os.unlink(doc)

if __name__ == '__main__' :

    copy = CopyPic()
    copy.rec_size(24)
    copy.rec_filepath('D:\Pic/')
    copy.excute_copy()
